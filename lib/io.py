"""
Orion IO Engine — Entrada y salida futurista.
────────────────────────────────────────────
Lectura universal de archivos (texto, PDF, Word, JSON, CSV),
extracción de información semántica,
escritura inteligente y streams masivos.
La consola y los archivos como flujos semánticos de datos.
"""

import os
import sys
import json
import chardet
import re
import unicodedata

from io import BytesIO
from core.types import OrionString

# ╭───────────────────────────────────────────────╮
# │ Consola: interacción viva con el entorno      │
# ╰───────────────────────────────────────────────╯

def io_show(*args, sep=" ", end="\n", env=None):
    """Imprime en consola estilo Orion con interpolación y formato amigable."""
    def orion_str(a):
        if env is not None and isinstance(a, OrionString):
            return str(a.interpolate(env))
        if isinstance(a, bool):
            return "yes" if a else "no"
        return str(a)
    sys.stdout.write(sep.join(orion_str(a) for a in args) + end)


def ask(prompt="> "):
    """Lee input del usuario."""
    return input(prompt)


# ╭───────────────────────────────────────────────╮
# │ Limpieza y normalización de texto             │
# ╰───────────────────────────────────────────────╯

def _clean_text(text):
    """Normaliza texto, elimina caracteres rotos o invisibles."""
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7EáéíóúÁÉÍÓÚñÑ]", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


# ╭───────────────────────────────────────────────╮
# │ Lectura universal                             │
# ╰───────────────────────────────────────────────╯

def orion_read(path):
    """
    Lee un archivo de cualquier tipo y devuelve un OrionString o estructura.
    Soporta: .txt, .log, .csv, .json, .pdf, .docx
    """
    ext = os.path.splitext(path)[1].lower()
    if ext in (".txt", ".csv", ".log"):
        return OrionString(_read_text_auto(path))
    if ext == ".json":
        return json.loads(_read_text_auto(path))
    if ext == ".pdf":
        return OrionString(_read_pdf(path))
    if ext in (".docx", ".doc"):
        return OrionString(_read_docx(path))
    raise ValueError(f"Tipo de archivo no soportado: {ext}")


def _read_text_auto(path):
    """Lee texto con detección automática de codificación."""
    with open(path, "rb") as f:
        raw = f.read()
    enc = chardet.detect(raw)["encoding"] or "utf-8"
    text = raw.decode(enc, errors="ignore")
    return _clean_text(text)


def _read_pdf(path):
    """Extrae texto de un PDF si pdfminer está disponible."""
    try:
        from pdfminer.high_level import extract_text
        return _clean_text(extract_text(path))
    except Exception as e:
        raise RuntimeError(f"Error leyendo PDF: {e}")


def _read_docx(path):
    """Extrae texto de un documento Word si python-docx está disponible."""
    try:
        import docx
        doc = docx.Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return _clean_text(text)
    except Exception as e:
        raise RuntimeError(f"Error leyendo Word: {e}")


# ╭───────────────────────────────────────────────╮
# │ Extracción inteligente Orion                  │
# ╰───────────────────────────────────────────────╯

def orion_extract(path, mode="auto"):
    """
    Extrae información estructurada de un archivo.
    - mode="auto": detecta y devuelve texto, números o tablas.
    - mode="tables": extrae tablas de PDF, CSV o texto tabular.
    - mode="numbers": extrae solo números o montos.
    - mode="emails": extrae correos electrónicos.
    - mode="dates": extrae fechas.
    """
    text = str(orion_read(path))

    if mode == "auto":
        return {
            "text": text[:5000] + ("..." if len(text) > 5000 else ""),
            "numbers": _extract_numbers(text),
            "emails": _extract_emails(text),
            "dates": _extract_dates(text)
        }

    if mode == "tables":
        return _extract_tables(path, text)
    if mode == "numbers":
        return _extract_numbers(text)
    if mode == "emails":
        return _extract_emails(text)
    if mode == "dates":
        return _extract_dates(text)

    raise ValueError(f"Modo de extracción no soportado: {mode}")


def _extract_numbers(text):
    """Extrae números, montos y posibles valores decimales."""
    return re.findall(r"-?\d+(?:[\.,]\d+)?", text)


def _extract_emails(text):
    """Extrae direcciones de correo electrónico."""
    return re.findall(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)


def _extract_dates(text):
    """Extrae fechas en varios formatos comunes."""
    patterns = [
        r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
        r"\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b",
        r"\b(?:ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)[a-z]*\s+\d{1,2},?\s+\d{2,4}\b"
    ]
    results = []
    for p in patterns:
        results += re.findall(p, text, flags=re.IGNORECASE)
    return results


def _extract_tables(path, text):
    """Intenta extraer tablas simples de CSV, PDF o texto plano."""
    ext = os.path.splitext(path)[1].lower()

    if ext == ".csv":
        import csv
        with open(path, encoding="utf-8") as f:
            reader = csv.reader(f)
            return [row for row in reader if any(row)]

    if ext == ".pdf":
        try:
            import pdfplumber
            tables = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    for table in page.extract_tables():
                        tables.append(table)
            return tables
        except Exception as e:
            raise RuntimeError(f"Error extrayendo tablas de PDF: {e}")

    # Tablas en texto: detecta columnas separadas por espacios o tabs
    lines = text.splitlines()
    tables = [re.split(r"\s{2,}|\t", l.strip()) for l in lines if re.search(r"\s{2,}|\t", l)]
    return [t for t in tables if len(t) > 1]


# ╭───────────────────────────────────────────────╮
# │ Escritura universal                           │
# ╰───────────────────────────────────────────────╯

def orion_write(path, data):
    """
    Escribe datos en un archivo detectando el formato:
    - .json → serializa
    - .txt / .log → texto plano
    - .docx → documento Word
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == ".json":
        _write_json(path, data)
    elif ext in (".txt", ".log", ".csv"):
        _write_text(path, data)
    elif ext in (".docx",):
        _write_docx(path, data)
    else:
        raise ValueError(f"Tipo de archivo no soportado para escritura: {ext}")


def _write_text(path, data, mode="w"):
    """Escribe texto plano (OrionString o str)."""
    if isinstance(data, (dict, list)):
        data = json.dumps(data, indent=2)
    with open(path, mode, encoding="utf-8") as f:
        f.write(str(data))


def _write_json(path, data):
    """Escribe JSON con indentación."""
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def _write_docx(path, text):
    """Escribe texto en un archivo Word (.docx)."""
    try:
        import docx
        doc = docx.Document()
        for line in str(text).split("\n"):
            doc.add_paragraph(line)
        doc.save(path)
    except Exception as e:
        raise RuntimeError(f"Error escribiendo Word: {e}")


# ╭───────────────────────────────────────────────╮
# │ Streams grandes (lazy)                        │
# ╰───────────────────────────────────────────────╯

def orion_stream(path):
    """
    Lee archivos de texto línea por línea, útil para archivos enormes.
    Devuelve líneas limpias (sin símbolos raros).
    """
    ext = os.path.splitext(path)[1].lower()
    if ext not in (".txt", ".log", ".csv"):
        raise ValueError("orion_stream solo admite archivos de texto")

    with open(path, "rb") as f:
        enc = chardet.detect(f.read(4096))["encoding"] or "utf-8"
        f.seek(0)
        for line in f:
            yield OrionString(_clean_text(line.decode(enc, errors="ignore")))


# ╭───────────────────────────────────────────────╮
# │ JSON helpers clásicos                         │
# ╰───────────────────────────────────────────────╯

def read_json(path):
    return json.loads(_read_text_auto(path))


def write_json(path, data):
    _write_json(path, data)


# ╭───────────────────────────────────────────────╮
# │ Accesos básicos para compatibilidad Orion     │
# ╰───────────────────────────────────────────────╯

def read_file(path, mode="r"):
    with open(path, mode, encoding="utf-8") as f:
        return f.read()


def write_file(path, content, mode="w"):
    with open(path, mode, encoding="utf-8") as f:
        f.write(content)


def append_file(path, content):
    write_file(path, content, mode="a")
